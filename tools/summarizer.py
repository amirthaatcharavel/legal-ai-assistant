from groq import Groq
import os
from dotenv import load_dotenv
import PyPDF2
from docx import Document
import io

# =========================
# 🔐 LOAD ENV
# =========================
load_dotenv()
client = Groq(api_key=os.getenv("GROQ_API_KEY"))

import re
import urllib.parse

def make_citations_clickable(text):
    if not text:
        return text

    pattern = r'(\d{4}\s+(AIR|SCC|SCR)[^,\n]*)'

    def replacer(match):
        citation = match.group(1)
        encoded = urllib.parse.quote(citation)

        link = f"https://indiankanoon.org/search/?formInput={encoded}"

        return f'<a href="{link}" target="_blank" style="color:#facc15;text-decoration:underline;">{citation}</a>'

    return re.sub(pattern, replacer, text)
# =========================
# 🔥 CHUNK TEXT
# =========================
def chunk_text(text, size=800):
    words = text.split()
    return [" ".join(words[i:i+size]) for i in range(0, len(words), size)]


# =========================
# 🔥 AI CASE SUMMARY (MAIN)
# =========================
def summarize(text, title):
    if not text:
        return f"""Case Name: {title}

Facts:
- Not available

Issue:
- Not available

Judgment:
- Not available

Reasoning:
- Not available
"""

    # 🔥 LIMIT INPUT (VERY IMPORTANT)
    text = text[:20000]

    chunks = chunk_text(text)
    partial_summaries = []

    # =========================
    # 🔥 STEP 1: SUMMARIZE CHUNKS
    # =========================
    for chunk in chunks[:5]:  # limit API calls

        prompt = f"""
You are a legal expert.

Summarize the case below STRICTLY in this format:

Facts:
- ...

Issue:
- ...

Judgment:
- ...

Reasoning:
- ...

Rules:
- Use bullet points ONLY
- Keep each point short
- Do NOT repeat content
- Do NOT add extra text

Case:
{chunk}
"""

        try:
            res = client.chat.completions.create(
                messages=[{"role": "user", "content": prompt}],
                model="llama-3.1-8b-instant"
            )

            partial_summaries.append(
                res.choices[0].message.content.strip()
            )

        except Exception as e:
            print("❌ Groq chunk error:", e)

    # =========================
    # 🔥 STEP 2: MERGE SUMMARIES
    # =========================
    final_prompt = f"""
You are a legal expert.

Merge these summaries into ONE final structured summary.

{partial_summaries}

Rules:
- Remove duplicates
- Keep only important points
- Keep it concise
- Maintain format EXACTLY

Return ONLY this format:

Facts:
- ...

Issue:
- ...

Judgment:
- ...

Reasoning:
- ...
"""

    try:
        final = client.chat.completions.create(
            messages=[{"role": "user", "content": final_prompt}],
            model="llama-3.1-8b-instant"
        )

        final_output = final.choices[0].message.content.strip()

    except Exception as e:
        print("❌ Groq final error:", e)
        final_output = "⚠️ Failed to generate final summary."

    # =========================
    # 🔥 FINAL OUTPUT
    # =========================
    return f"Case Name: {title}\n\n{final_output}"


# =========================
# 🔥 SIMPLE TEXT SUMMARY (FOR FILE UPLOAD)
# =========================
def summarize_text(text):
    if not text:
        return "No content to summarize."

    import re

    # =========================
    # 🔥 STEP 1 — LIMIT SIZE
    # =========================
    text = text[:15000]

    # =========================
    # 🔥 STEP 2 — REMOVE LEGAL GARBAGE
    # =========================
    patterns = [
        r"IA\s*No\.\s*\d+\/\d+.*",
        r"ON IA.*",
        r"FOR EXEMPTION.*",
        r"PERMISSION TO FILE.*",
        r"APPLICATION FOR.*",
        r"APPROPRIATE ORDERS.*",
        r"W\.P.*No\..*",
        r"SLP.*No\..*",
        r"CRL.*No\..*",
    ]

    for p in patterns:
        text = re.sub(p, "", text, flags=re.IGNORECASE)

    # =========================
    # 🔥 STEP 3 — REMOVE METADATA LINES
    # =========================
    lines = text.split("\n")
    clean_lines = []

    skip_words = [
        "equivalent citations",
        "author",
        "bench",
        "petitioner",
        "respondent",
        "advocate",
        "appearance",
    ]

    for line in lines:
        if not any(word in line.lower() for word in skip_words):
            clean_lines.append(line)

    clean_text = " ".join(clean_lines)

    # =========================
    # 🔥 STEP 4 — SPLIT SENTENCES
    # =========================
    sentences = re.split(r'(?<=[.!?]) +', clean_text)

    # 🔥 Remove very short / useless sentences
    sentences = [s.strip() for s in sentences if len(s.strip()) > 30]

    # 🔥 Skip initial garbage (very important)
    sentences = sentences[5:60]

    # =========================
    # 🧠 STEP 5 — CLASSIFY
    # =========================
    facts = []
    issue = []
    judgment = []
    reasoning = []

    for s in sentences:
        s_lower = s.lower()

        if any(word in s_lower for word in ["fact", "background", "arose", "case of"]):
            facts.append(s)

        elif any(word in s_lower for word in ["issue", "question", "whether"]):
            issue.append(s)

        elif any(word in s_lower for word in ["held", "court held", "decided", "judgment"]):
            judgment.append(s)

        elif any(word in s_lower for word in ["reason", "because", "analysis", "observed"]):
            reasoning.append(s)

    # =========================
    # 🔁 STEP 6 — FALLBACKS
    # =========================
    if not facts:
        facts = sentences[:3]

    if not issue:
        issue = sentences[3:6]

    if not judgment:
        judgment = sentences[6:9]

    if not reasoning:
        reasoning = sentences[9:12]

    # =========================
    # 🎯 STEP 7 — FINAL OUTPUT
    # =========================
    summary = f"""
Facts:
{" ".join(facts[:3])}

Issue:
{" ".join(issue[:2])}

Judgment:
{" ".join(judgment[:2])}

Reasoning:
{" ".join(reasoning[:3])}
"""

    summary = make_citations_clickable(summary)
    return summary.strip()
# =========================
# 🔥 FILE TEXT EXTRACTION
# =========================
def extract_text(filename, content):

    try:
        # PDF
        if filename.endswith(".pdf"):
            pdf = PyPDF2.PdfReader(io.BytesIO(content))
            text = ""

            for page in pdf.pages:
                text += page.extract_text() or ""

            return text

        # DOCX
        elif filename.endswith(".docx"):
            doc = Document(io.BytesIO(content))
            return "\n".join([p.text for p in doc.paragraphs])

        # TXT
        elif filename.endswith(".txt"):
            return content.decode("utf-8", errors="ignore")

        else:
            return "Unsupported file format"

    except Exception as e:
        print("❌ Extraction error:", e)
        return ""

import pytesseract
from PIL import Image
import io

import io
from PIL import Image
import pytesseract

def extract_text(filename, content):
    try:
        # 🔥 FIX: handle uppercase extensions
        filename = filename.lower()

        # =========================
        # 📄 PDF FILE
        # =========================
        if filename.endswith(".pdf"):
            import PyPDF2

            pdf = PyPDF2.PdfReader(io.BytesIO(content))
            text = ""

            for page in pdf.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted + "\n"

            # 🔥 fallback if PDF has no text (scanned)
            if not text.strip():
                return "⚠️ No readable text found (scanned PDF). Use OCR."

            return text

        # =========================
        # 📄 DOCX FILE
        # =========================
        elif filename.endswith(".docx"):
            from docx import Document

            doc = Document(io.BytesIO(content))
            return "\n".join([p.text for p in doc.paragraphs])

        # =========================
        # 📄 TXT FILE
        # =========================
        elif filename.endswith(".txt"):
            return content.decode("utf-8", errors="ignore")

        # =========================
        # 🖼 IMAGE FILE (OCR)
        # =========================
        elif filename.endswith((".png", ".jpg", ".jpeg")):
            image = Image.open(io.BytesIO(content))
            text = pytesseract.image_to_string(image)

            if not text.strip():
                return "⚠️ No text detected in image."

            return text

        # =========================
        # ❌ UNSUPPORTED
        # =========================
        else:
            return "Unsupported file format"

    except Exception as e:
        print("❌ Extraction Error:", e)
        return "⚠️ Error extracting text"

